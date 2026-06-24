#!/usr/bin/env python3
"""One-page Canadian-style resume with summary and quantified bullets."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "Yang_Zhang_Resume.docx"

BODY_SIZE = Pt(9)
SECTION_SIZE = Pt(10)
NAME_SIZE = Pt(19)
SUBHEADER_SIZE = Pt(8.5)


def set_run_font(run, size=None, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text.upper())
    set_run_font(run, SECTION_SIZE, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    pBdr.append(bottom)


def add_compact_line(doc: Document, text: str, space_after: float = 1) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        set_run_font(run, BODY_SIZE)


def add_project_header(doc: Document, title: str, role_dates: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(title)
    set_run_font(r1, BODY_SIZE, bold=True)
    r2 = p.add_run(f"  |  {role_dates}")
    set_run_font(r2, BODY_SIZE)


def add_job_header(doc: Document, title: str, org: str, dates: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(title)
    set_run_font(r1, BODY_SIZE, bold=True)
    r2 = p.add_run(f"  |  {org}  |  {dates}")
    set_run_font(r2, BODY_SIZE)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            set_run_font(run, BODY_SIZE)


def build_resume() -> Document:
    doc = Document()
    set_document_style(doc)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.4)
    sec.bottom_margin = Inches(0.4)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    name_p = doc.add_paragraph()
    set_run_font(name_p.add_run("YANG ZHANG"), NAME_SIZE, bold=True, name="Georgia")

    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(2)
    set_run_font(
        contact_p.add_run(
            "Burnaby, BC  |  sc20190702@gmail.com  |  "
            "linkedin.com/in/yang-zhang-527094327  |  github.com/Yang-Zhang1994"
        ),
        SUBHEADER_SIZE,
        color=RGBColor(0x44, 0x44, 0x44),
    )

    add_section_title(doc, "Professional Summary")
    add_compact_line(
        doc,
        "MSCS candidate (GPA 4.0/4.0) specializing in enterprise-grade distributed systems and "
        "cloud CI/CD. Shipped a 12-service e-commerce platform on AWS EKS with a live HTTPS "
        "storefront, admin portal, and Stripe checkout. Proficient in Java/Spring Boot "
        "microservices, React/Next.js, PostgreSQL, and Kubernetes delivery (Terraform, Helm, "
        "GitHub Actions). Background in data-driven neuroscience research and cross-functional "
        "stakeholder collaboration. Bilingual: English and Mandarin.",
    )

    add_section_title(doc, "Technical Skills")
    add_compact_line(
        doc,
        "Languages: Java, Python, TypeScript, JavaScript, SQL, C",
    )
    add_compact_line(
        doc,
        "Backend & Microservices: Spring Boot 3, Spring Cloud Gateway, Consul, Node.js, Express, "
        "Prisma, RESTful APIs",
    )
    add_compact_line(doc, "Frontend: React 19, Next.js, Vue, HTML5/CSS3")
    add_compact_line(
        doc,
        "Databases & Caching: PostgreSQL, Redis (multi-level caching), Elasticsearch, MongoDB, "
        "Amazon RDS",
    )
    add_compact_line(
        doc,
        "Cloud & DevOps: AWS (EKS, ECR, ALB, S3, RDS), Terraform (IaC), Helm, Kubernetes, Docker, "
        "GitHub Actions, Linux",
    )
    add_compact_line(
        doc,
        "Testing & Tools: JUnit, Mockito, Playwright, Vitest, k6 (load testing), Git, Postman, "
        "Stripe API",
        space_after=0,
    )

    add_section_title(doc, "Software Engineering Projects")

    add_project_header(
        doc,
        "GrainMart — Distributed E-Commerce Platform (Microservices)",
        "Lead Full-Stack Developer  |  2025 – Present",
    )
    add_bullets(
        doc,
        [
            "Architected 12 Spring Boot microservices with database-per-service on AWS RDS "
            "PostgreSQL, Consul service discovery, and Spring Cloud Gateway routing.",
            "Offloaded product search to Elasticsearch; built Redis cache-aside with guards "
            "against cache penetration, breakdown, and avalanche on read-heavy catalog paths.",
            "Delivered Next.js storefront and Vue admin console with Stripe Checkout, Google "
            "OAuth2, RabbitMQ async orders, and S3 presigned uploads; idempotent orders via "
            "unique constraints and Stripe webhook signature verification.",
            "Deployed on AWS EKS (Terraform, Helm, ALB + ACM HTTPS) across 3 domains; GitHub "
            "Actions CI/CD with OIDC (9 JUnit gate, SHA-tagged ECR images, Helm deploy, GitHub "
            "Secrets); Gateway/Product CPU HPA; k6 load test ~36 RPS on product search; 2 "
            "Playwright E2E suites and 10-step API purchase validation.",
        ],
    )
    add_compact_line(
        doc,
        "Repository: github.com/Yang-Zhang1994/ecommerce-microservices-backend  |  "
        "Live: mall.yangzhangtech.online  |  Admin: admin.yangzhangtech.online (demo login in README)",
        space_after=0,
    )

    add_project_header(
        doc,
        "SmartChef — Recipe & Ingredient-Matching Web App",
        "Full-Stack Developer  |  2025",
    )
    add_bullets(
        doc,
        [
            "Built React 19 + Node/Express + Prisma/PostgreSQL app with ingredient-matching search "
            "and role-based access; seeded catalog (60+ ingredients, 30 recipes); 8 Vitest "
            "component tests.",
            "Deployed on Vercel + Render with bcrypt auth and HTTP-only JWT cookies "
            "(final-project-kangaroo.vercel.app).",
        ],
    )

    add_project_header(
        doc,
        "BC PhysEd Educational Tool (EdTech Platform)",
        "Backend Developer  |  2025",
    )
    add_bullets(
        doc,
        [
            "Designed REST APIs and a teacher reporting portal for a Grades 4–7 LMS.",
            "Built a reward-scaling algorithm driven by real-time student performance analytics "
            "(github.com/Yang-Zhang1994/bc-physed-digital-learning-tool).",
        ],
    )

    add_section_title(doc, "Professional & Research Experience")
    add_job_header(
        doc,
        "Research Assistant (Data-Driven Development)",
        "East China Normal University",
        "2020 – 2023",
    )
    add_bullets(
        doc,
        [
            "Translated neuroscience research into software requirements with K-12 educators; "
            "managed IRB protocols and coordinated national workshops to inform user-centred "
            "EdTech API design.",
            "Led 4 experimental studies; analyzed behavioural datasets with statistical methods "
            "to inform product and learning-design decisions.",
        ],
    )

    add_section_title(doc, "Education")
    add_compact_line(
        doc,
        "Northeastern University — M.S. in Computer Science (GPA: 4.0/4.0)  |  2024 – Present",
    )
    add_compact_line(
        doc,
        "East China Normal University — M.Ed. in Applied Psychology (GPA: 3.61/4.0)  |  2018 – 2021",
    )
    add_compact_line(
        doc,
        "Central China Normal University — B.Sc. in Chemistry  |  2013 – 2017",
        space_after=0,
    )

    return doc


def main() -> None:
    doc = build_resume()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
